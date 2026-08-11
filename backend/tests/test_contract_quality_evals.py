"""Self-tests and the hard gate for the deterministic quality corpus."""

from __future__ import annotations

import copy
import json
import socket
import subprocess
import sys
from dataclasses import replace
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import pytest
from bs4 import BeautifulSoup
from docx import Document

import quality_evals.__main__ as quality_cli
import quality_evals.offline_gate as offline_gate
import quality_evals.offline_worker as offline_worker
import quality_evals.runner as quality_runner
from app.draft_state import required_field_keys, snapshot_from_document_state
from app.export import build_export_document, render_docx, render_pdf
from app.manifests import load_manifest
from quality_evals.corpus import (
    CorpusDocument,
    CorpusValidationError,
    ValidatedCorpus,
    load_corpus,
    validate_corpus,
)
from quality_evals.report import (
    HARD_GATE_METRICS,
    METRIC_NAMES,
    CoverageExpectation,
    QualityCaseResult,
    QualityReport,
)
from quality_evals.runner import run_contract_quality_evaluation

EXPECTED_METRIC_NAMES = {
    "unknown_field_acceptance_count",
    "invalid_type_acceptance_count",
    "required_field_false_negative_count",
    "required_field_false_positive_count",
    "conflict_transition_failure_count",
    "successful_invalid_downloads",
    "cross_format_semantic_mismatch_count",
}
EXPECTED_HARD_GATE_METRICS = set(EXPECTED_METRIC_NAMES)


def test_quality_corpus_discovers_every_catalog_manifest():
    corpus = load_corpus()

    validated = validate_corpus(corpus)

    assert len(validated.documents) == 11
    assert validated.catalog_doc_ids == validated.manifest_doc_ids
    assert validated.catalog_doc_ids == validated.corpus_doc_ids


def test_quality_corpus_rejects_catalog_manifest_or_corpus_drift():
    corpus = load_corpus()

    with _raises_corpus_error("catalog_manifest_mismatch"):
        validate_corpus(corpus, catalog_doc_ids={*corpus.doc_ids, "new-document"})

    missing_document = copy.deepcopy(corpus.raw)
    missing_document["documents"] = missing_document["documents"][:-1]
    with _raises_corpus_error("catalog_corpus_mismatch"):
        validate_corpus(missing_document)


def test_quality_corpus_rejects_duplicate_case_ids_and_unknown_fields():
    corpus = load_corpus()
    duplicate = copy.deepcopy(corpus.raw)
    duplicate["cases"].append(copy.deepcopy(duplicate["cases"][0]))
    with _raises_corpus_error("duplicate_case_id"):
        validate_corpus(duplicate)

    unknown_field = copy.deepcopy(corpus.raw)
    unknown_field["documents"][0]["anchor_field"] = "不存在字段"
    with _raises_corpus_error("unknown_corpus_field"):
        validate_corpus(unknown_field)


def test_quality_corpus_rejects_missing_case_suite():
    corpus = load_corpus()
    incomplete = copy.deepcopy(corpus.raw)
    incomplete["cases"] = incomplete["cases"][:-1]

    with _raises_corpus_error("case_coverage_mismatch"):
        validate_corpus(incomplete)


def test_quality_corpus_rejects_unregistered_case_kind():
    corpus = load_corpus()
    changed = copy.deepcopy(corpus.raw)
    changed["cases"][0]["kind"] = "not_registered"

    with _raises_corpus_error("unregistered_case_kind"):
        validate_corpus(
            changed,
            registered_case_kinds=set(quality_runner.CASE_REGISTRY),
        )


def test_quality_corpus_requires_both_renderers():
    corpus = load_corpus()
    missing_pdf = copy.deepcopy(corpus.raw)
    missing_pdf["renderers"] = ["docx"]

    with _raises_corpus_error("renderer_coverage_missing"):
        validate_corpus(missing_pdf)


@pytest.mark.parametrize(
    ("raw", "kind"),
    [
        ([], "invalid_corpus"),
        (None, "invalid_corpus"),
        ("not-an-object", "invalid_corpus"),
        ({"schema_version": True}, "unsupported_corpus_schema"),
        (
            {
                "schema_version": 1,
                "renderers": ["docx", "pdf"],
                "cases": [{"id": [], "kind": "complete_state"}],
                "documents": [],
            },
            "invalid_corpus",
        ),
        (
            {
                "schema_version": 1,
                "renderers": ["docx", "pdf"],
                "cases": [{"id": "complete", "kind": {}}],
                "documents": [],
            },
            "invalid_corpus",
        ),
        (
            {
                "schema_version": 1,
                "renderers": ["docx", "pdf"],
                "cases": [],
                "documents": [{"doc_id": [], "anchor_field": "字段"}],
            },
            "invalid_corpus",
        ),
        (
            {
                "schema_version": 1,
                "renderers": ["docx", "pdf"],
                "cases": [],
                "documents": [{"doc_id": "doc", "anchor_field": {}}],
            },
            "invalid_corpus",
        ),
        (
            {
                "schema_version": 1,
                "renderers": ["docx", 1],
                "cases": [],
                "documents": [],
            },
            "invalid_corpus",
        ),
    ],
)
def test_quality_corpus_rejects_malformed_shapes(raw, kind: str):
    with _raises_corpus_error(kind):
        validate_corpus(raw)


def test_quality_corpus_honors_explicit_empty_document_sets():
    with _raises_corpus_error("catalog_corpus_mismatch"):
        validate_corpus(
            load_corpus(),
            catalog_doc_ids=set(),
            manifest_doc_ids=set(),
        )


def test_condition_evaluator_exercises_required_when_arrays_as_conjunctions():
    manifest = {
        "doc_id": "synthetic-conjunction",
        "version": 1,
        "fields": [
            {
                "key": "模式",
                "type": "string",
                "enum": ["付费", "免费"],
                "required": False,
            },
            {
                "key": "地区",
                "type": "string",
                "options": ["境内", "境外"],
                "required": False,
            },
            {
                "key": "日期",
                "type": "date",
                "required": False,
            },
            {
                "key": "付款安排",
                "type": "text",
                "required": False,
                "required_when": [
                    {"field": "模式", "op": "equals", "value": "付费"},
                    {"field": "地区", "op": "in", "values": ["境内"]},
                    {
                        "field": "日期",
                        "op": "not_equals",
                        "value": "2026-01-16",
                    },
                ],
            },
        ],
    }
    report = QualityReport.empty(expected_doc_ids={manifest["doc_id"]})

    quality_runner._evaluate_conditions(
        report,
        CorpusDocument(doc_id=manifest["doc_id"], anchor_field="模式"),
        manifest,
    )

    assert report.failed_cases == 0, report.failure_summary()


def test_condition_evaluator_rejects_impossible_conjunction_witnesses():
    manifest = {
        "doc_id": "impossible-conjunction",
        "version": 1,
        "fields": [
            {
                "key": "模式",
                "type": "string",
                "enum": ["甲", "乙"],
                "required": False,
            },
            {
                "key": "结果",
                "type": "string",
                "required": False,
                "required_when": [
                    {"field": "模式", "op": "equals", "value": "甲"},
                    {"field": "模式", "op": "equals", "value": "乙"},
                ],
            },
        ],
    }

    with _raises_corpus_error("condition_witness_unavailable"):
        quality_runner._evaluate_conditions(
            QualityReport.empty(expected_doc_ids={manifest["doc_id"]}),
            CorpusDocument(doc_id=manifest["doc_id"], anchor_field="模式"),
            manifest,
        )


@pytest.mark.parametrize(
    "conditions",
    [
        [
            {"field": "模式", "op": "equals", "value": "甲"},
            {"field": "模式", "op": "in", "values": ["甲", "乙"]},
        ],
        [
            {"field": "模式", "op": "equals", "value": "甲"},
            {"field": "模式", "op": "equals", "value": "甲"},
        ],
    ],
)
def test_condition_evaluator_accepts_feasible_same_driver_conjunctions(conditions):
    manifest = {
        "doc_id": "same-driver-conjunction",
        "version": 1,
        "fields": [
            {
                "key": "模式",
                "type": "string",
                "enum": ["甲", "乙"],
                "required": False,
            },
            {
                "key": "结果",
                "type": "string",
                "required": False,
                "required_when": conditions,
            },
        ],
    }
    report = QualityReport.empty(expected_doc_ids={manifest["doc_id"]})

    quality_runner._evaluate_conditions(
        report,
        CorpusDocument(doc_id=manifest["doc_id"], anchor_field="模式"),
        manifest,
    )

    assert report.failed_cases == 0, report.failure_summary()


def test_shared_condition_conformance_vectors_match_backend_semantics():
    raw = json.loads(
        Path(quality_runner.__file__)
        .with_name("condition_conformance.json")
        .read_text()
    )
    assert raw["schema_version"] == 1
    for vector in raw["vectors"]:
        manifest = vector.get("manifest") or load_manifest(
            vector["manifest_doc_id"]
        )
        assert manifest is not None
        conditional_keys = {
            field["key"]
            for field in manifest["fields"]
            if field.get("required_when") is not None
        }
        if vector["positive_witness"] == "impossible":
            dependent = next(
                field
                for field in manifest["fields"]
                if field.get("required_when") is not None
            )
            conditions = dependent["required_when"]
            conditions = conditions if isinstance(conditions, list) else [conditions]
            with _raises_corpus_error("condition_witness_unavailable"):
                quality_runner._condition_group_witnesses(manifest, conditions)
            continue
        for assignment in vector["assignments"]:
            snapshot = snapshot_from_document_state(
                doc_id=manifest["doc_id"], state={}, manifest=manifest
            )
            for key, value in assignment["values"].items():
                snapshot.fields[key].status = "confirmed"
                snapshot.fields[key].value = value
                snapshot.fields[key].confirmed_at = "2026-01-15T00:00:00+00:00"
                snapshot.fields[key].confirmed_by_user_id = 1
            for key in assignment.get("unconfirmed", []):
                snapshot.fields[key].status = "pending_confirmation"
                snapshot.fields[key].confirmed_at = None
                snapshot.fields[key].confirmed_by_user_id = None
            actual = set(required_field_keys(manifest, snapshot)) & conditional_keys
            assert actual == set(assignment["expected_active"]), vector["id"]


def test_nonisolatable_group_negative_is_not_a_corpus_failure():
    raw = json.loads(
        Path(quality_runner.__file__)
        .with_name("condition_conformance.json")
        .read_text()
    )
    vector = next(
        item for item in raw["vectors"] if item["id"] == "nonisolatable-negative"
    )
    manifest = vector["manifest"]
    dependent = manifest["fields"][1]

    positive, negatives = quality_runner._condition_group_witnesses(
        manifest, [dependent["required_when"]]
    )

    assert positive == {"x": "A"}
    assert negatives == {}


def test_condition_oracle_detects_collateral_activation(monkeypatch):
    manifest = {
        "doc_id": "synthetic-collateral",
        "version": 1,
        "fields": [
            {
                "key": "模式",
                "type": "string",
                "enum": ["启用", "停用"],
                "required": False,
            },
            {
                "key": "目标",
                "type": "string",
                "required": False,
                "required_when": {
                    "field": "模式",
                    "op": "equals",
                    "value": "启用",
                },
            },
            {
                "key": "旁路字段",
                "type": "string",
                "required": False,
                "required_when": {
                    "field": "模式",
                    "op": "equals",
                    "value": "停用",
                },
            },
        ],
    }
    original = quality_runner.required_field_keys

    def activate_collateral(manifest, snapshot):
        return [*original(manifest, snapshot), "旁路字段"]

    monkeypatch.setattr(quality_runner, "required_field_keys", activate_collateral)
    report = QualityReport.empty(expected_doc_ids={manifest["doc_id"]})
    document = CorpusDocument(manifest["doc_id"], "模式")

    quality_runner._record_condition_assignment(
        report,
        document,
        manifest,
        "required-when:目标:mutation",
        {"模式": "启用"},
        metric="required_field_false_positive_count",
    )

    assert report.failed_cases == 1
    assert "旁路字段" in report.failure_summary()


def test_hard_gate_report_exits_nonzero_for_any_gated_metric():
    assert set(METRIC_NAMES) == EXPECTED_METRIC_NAMES
    assert set(HARD_GATE_METRICS) == EXPECTED_HARD_GATE_METRICS
    for metric in EXPECTED_HARD_GATE_METRICS:
        report = _valid_report()
        assert report.exit_code == 0
        report.metrics[metric] = 1
        assert report.exit_code == 1


def test_report_hard_gate_enforces_denominators_schema_coverage_and_metric_shape():
    report = QualityReport.empty(expected_doc_ids={"expected-doc"})
    assert report.exit_code == 1
    assert "zero_metric_denominator" in report.invariant_errors
    assert "document_coverage_incomplete" in report.invariant_errors

    for metric in METRIC_NAMES:
        report.add(
            QualityCaseResult(
                doc_id="expected-doc",
                case_id=f"case-{metric}",
                passed=True,
                metric=metric,
            )
        )
    assert report.exit_code == 0

    report.schema_version = 2
    assert report.exit_code == 1
    assert "unsupported_report_schema" in report.invariant_errors

    report.schema_version = 1
    report.metrics["unregistered_metric"] = 0
    assert report.exit_code == 1
    assert "metric_shape_mismatch" in report.invariant_errors


def test_report_serializes_expected_actual_missing_and_duplicate_coverage():
    expectation = CoverageExpectation(
        doc_id="expected-doc",
        case_kind="complete_state",
        scenario="complete-state",
    )
    report = _valid_report()
    report.expected_coverage = (expectation,)
    report.add(
        QualityCaseResult(
            doc_id="expected-doc",
            case_id="coverage-one",
            passed=True,
            coverage_key=expectation.key,
        )
    )
    report.add(
        QualityCaseResult(
            doc_id="expected-doc",
            case_id="coverage-duplicate",
            passed=True,
            coverage_key=expectation.key,
        )
    )

    coverage = report.to_dict()["coverage"]
    assert coverage["expected_count"] == 1
    assert coverage["actual_count"] == 2
    assert coverage["missing_coverage_keys"] == []
    assert coverage["unexpected_duplicate_coverage_keys"] == [expectation.key]
    assert "coverage_duplicate" in report.invariant_errors


def test_report_fails_when_expected_coverage_is_missing():
    expectation = CoverageExpectation(
        doc_id="expected-doc",
        case_kind="complete_state",
        scenario="complete-state",
    )
    report = _valid_report()
    report.expected_coverage = (expectation,)

    assert report.to_dict()["coverage"]["missing_coverage_keys"] == [expectation.key]
    assert "coverage_missing" in report.invariant_errors


def test_quality_command_returns_nonzero_when_a_hard_gate_triggers(monkeypatch):
    report = _valid_report()
    assert report.exit_code == 0
    report.metrics[next(iter(EXPECTED_HARD_GATE_METRICS))] = 1
    monkeypatch.setattr(quality_cli, "run_contract_quality_evaluation", lambda: report)
    monkeypatch.setattr(sys, "argv", ["quality-evals"])

    assert quality_cli.main() == 1


def test_quality_json_command_returns_nonzero_for_serialization_failure(
    monkeypatch,
    capsys,
):
    report = QualityReport.empty()
    monkeypatch.setattr(report, "to_dict", lambda: {"bad": object()})
    monkeypatch.setattr(quality_cli, "run_contract_quality_evaluation", lambda: report)
    monkeypatch.setattr(sys, "argv", ["quality-evals", "--json"])

    assert quality_cli.main() == 1
    assert "report_serialization_failed" in capsys.readouterr().err


def test_quality_command_returns_stable_corpus_failure(monkeypatch, capsys):
    def fail_corpus():
        raise CorpusValidationError("synthetic_corpus_failure", "bad corpus")

    monkeypatch.setattr(quality_cli, "run_contract_quality_evaluation", fail_corpus)
    monkeypatch.setattr(sys, "argv", ["quality-evals", "--json"])

    assert quality_cli.main() == 1
    stderr = capsys.readouterr().err
    assert "corpus_validation_failed:synthetic_corpus_failure" in stderr


@pytest.mark.parametrize(
    ("probe", "marker"),
    [
        ("socket", "offline_guard_blocked_create_connection"),
        ("inet_socket", "offline_guard_blocked_socket_family:AF_INET"),
        ("low_level_socket", "offline_guard_blocked_socket_family:AF_INET"),
        ("ipv6", "offline_guard_blocked_socket_family:AF_INET6"),
        ("dns", "offline_guard_blocked_dns"),
        ("subprocess", "offline_guard_blocked_subprocess"),
        ("ldconfig_args", "offline_guard_blocked_subprocess"),
        ("app_llm", "offline_guard_blocked_import:app.llm"),
        ("litellm", "offline_guard_blocked_import:litellm"),
    ],
)
def test_fresh_process_offline_guard_rejects_network_and_llm_imports(
    probe: str, marker: str
):
    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "quality_evals.offline_gate",
            "--probe",
            probe,
        ],
        capture_output=True,
        text=True,
        check=False,
    )

    assert completed.returncode != 0
    assert marker in completed.stderr


def test_offline_guard_local_subprocess_allowlist_is_exact():
    args = (["/sbin/ldconfig", "-p"],)
    kwargs = {
        "stdin": subprocess.DEVNULL,
        "stderr": subprocess.DEVNULL,
        "stdout": subprocess.PIPE,
        "env": {"LC_ALL": "C", "LANG": "C"},
    }

    assert offline_worker._is_allowed_local_process(args, kwargs)
    assert not offline_worker._is_allowed_local_process(
        (["/sbin/ldconfig", "-p", "--extra"],), kwargs
    )
    assert not offline_worker._is_allowed_local_process(
        args, {**kwargs, "shell": True}
    )
    assert not offline_worker._is_allowed_local_process(
        args, {**kwargs, "env": {**kwargs["env"], "EXTRA": "1"}}
    )


def test_fresh_process_offline_guard_disables_ctypes_toolchain_fallbacks():
    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "quality_evals.offline_gate",
            "--probe",
            "ctypes_toolchain",
        ],
        capture_output=True,
        text=True,
        check=False,
    )

    assert completed.returncode == 0, completed.stderr
    assert "offline_guard_blocked" not in completed.stderr


@pytest.mark.skipif(not hasattr(socket, "AF_UNIX"), reason="AF_UNIX unavailable")
def test_fresh_process_offline_guard_allows_unix_socket():
    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "quality_evals.offline_gate",
            "--probe",
            "unix_socket",
        ],
        capture_output=True,
        text=True,
        check=False,
    )

    assert completed.returncode == 0, completed.stderr
    assert "offline_guard_blocked" not in completed.stderr


def test_fresh_process_udp_probe_closes_legacy_three_patch_escape():
    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "quality_evals.offline_gate",
            "--probe",
            "udp",
        ],
        capture_output=True,
        text=True,
        check=False,
    )

    assert completed.returncode != 0
    assert "offline_guard_blocked_socket_family:AF_INET" in completed.stderr
    assert "udp_bytes_sent=1" not in completed.stdout


def test_offline_gate_parses_json_and_preserves_evaluator_exit(monkeypatch):
    monkeypatch.setattr(
        offline_gate.subprocess,
        "run",
        lambda *_args, **_kwargs: SimpleNamespace(
            stdout='{"schema_version": 1}', stderr="", returncode=7
        ),
    )

    assert offline_gate.main() == 7


@pytest.mark.parametrize("status_code", [422, 500])
def test_invalid_download_conflict_setup_failure_is_a_hard_failure(
    monkeypatch,
    status_code: int,
):
    original_patch = quality_runner._api_patch

    def fail_conflict_patch(*args, **kwargs):
        if "download-conflict" in kwargs.get("patch_id", ""):
            return SimpleNamespace(
                status_code=status_code,
                text="injected conflict setup failure",
                json=lambda: {},
            )
        return original_patch(*args, **kwargs)

    monkeypatch.setattr(quality_runner, "_api_patch", fail_conflict_patch)
    document = CorpusDocument(doc_id="mutual-nda", anchor_field="保密用途")
    manifest = load_manifest(document.doc_id)
    assert manifest is not None
    report = QualityReport.empty(expected_doc_ids={document.doc_id})

    with quality_runner._evaluation_client() as (client, get_conn):
        headers = quality_runner._register(client)
        quality_runner._evaluate_invalid_downloads(
            report,
            client,
            get_conn,
            headers,
            document,
            manifest,
        )

    failures = [result for result in report.results if not result.passed]
    assert any(
        result.case_id == "invalid-downloads:conflict-setup"
        and f"HTTP {status_code}" in (result.actual or "")
        for result in failures
    )
    assert report.exit_code == 1
    assert "conflict-setup" in report.failure_summary()


def test_semantic_verifier_rejects_empty_wrong_and_corrupt_downloads():
    model, docx_payload, pdf_payload = _semantic_fixture()
    valid_headers = {
        "content-type": (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        "content-disposition": "attachment; filename*=UTF-8''quality.docx",
    }

    assert quality_runner._semantic_response_issues(
        "docx",
        _download_response(docx_payload, valid_headers),
        model,
        forbidden_values=(),
    ) == []
    assert quality_runner._semantic_response_issues(
        "docx",
        _download_response(b"", valid_headers),
        model,
        forbidden_values=(),
    )
    assert quality_runner._semantic_response_issues(
        "docx",
        _download_response(pdf_payload, valid_headers),
        model,
        forbidden_values=(),
    )
    assert quality_runner._semantic_response_issues(
        "docx",
        _download_response(b"PKbroken", valid_headers),
        model,
        forbidden_values=(),
    )
    assert quality_runner._semantic_response_issues(
        "pdf",
        _download_response(
            b"%PDF-broken",
            {
                "content-type": "application/pdf",
                "content-disposition": "attachment; filename*=UTF-8''quality.pdf",
            },
        ),
        model,
        forbidden_values=(),
    )
    assert quality_runner._semantic_response_issues(
        "docx",
        _download_response(docx_payload, {}),
        model,
        forbidden_values=(),
    )


def test_semantic_verifier_requires_title_fields_terms_disclaimer_and_every_block():
    model, _docx_payload, _pdf_payload = _semantic_fixture()
    headers = {
        "content-type": (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        "content-disposition": "attachment; filename*=UTF-8''quality.docx",
    }
    values = [field.value for field in model.fields if field.value]
    only_fields = _docx_with_lines([model.title, *values, model.disclaimer])
    only_fields_issues = quality_runner._semantic_response_issues(
        "docx",
        _download_response(only_fields, headers),
        model,
        forbidden_values=(),
    )
    assert any("standard terms" in issue for issue in only_fields_issues)

    omitted = next(block for block in model.blocks if block.section == "standard_terms")
    without_one_block = _docx_with_lines(
        [
            model.title,
            *values,
            *(block.text for block in model.blocks if block != omitted),
            model.disclaimer,
        ]
    )
    block_issues = quality_runner._semantic_response_issues(
        "docx",
        _download_response(without_one_block, headers),
        model,
        forbidden_values=(),
    )
    assert any("block" in issue for issue in block_issues)


def test_semantic_verifier_rejects_payload_from_the_wrong_snapshot():
    model, _docx_payload, _pdf_payload = _semantic_fixture(anchor_value="预期稳定值")
    wrong_model, wrong_payload, _pdf = _semantic_fixture(anchor_value="错误快照值")
    assert wrong_model != model
    issues = quality_runner._semantic_response_issues(
        "docx",
        _download_response(
            wrong_payload,
            {
                "content-type": (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                "content-disposition": "attachment; filename*=UTF-8''quality.docx",
            },
        ),
        model,
        forbidden_values=(),
    )
    assert any("field" in issue for issue in issues)


@pytest.mark.parametrize(
    ("renderer", "state_name"),
    [("docx", "conflict"), ("pdf", "conflict"), ("docx", "pending")],
)
def test_semantic_verifier_detects_adapter_candidate_leakage(
    renderer: str, state_name: str
):
    model, _docx_payload, _pdf_payload = _semantic_fixture()
    candidate = f"PL24-{state_name}-adapter-leak"
    leaked = replace(
        model,
        html=model.html.replace(
            "</section>\n</main>", f"<p>{candidate}</p></section>\n</main>"
        ),
    )
    payload = render_docx(leaked) if renderer == "docx" else render_pdf(leaked)
    content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if renderer == "docx"
        else "application/pdf"
    )
    issues = quality_runner._semantic_response_issues(
        renderer,
        _download_response(
            payload,
            {
                "content-type": content_type,
                "content-disposition": (
                    f"attachment; filename*=UTF-8''quality.{renderer}"
                ),
            },
        ),
        model,
        forbidden_values=(candidate,),
    )

    assert any("candidate leaked" in issue for issue in issues)


def test_semantic_verifier_does_not_flag_authoritative_value_substrings():
    model, docx_payload, _pdf_payload = _semantic_fixture()
    stable_value = next(field.value for field in model.fields if field.value)

    issues = quality_runner._semantic_response_issues(
        "docx",
        _download_response(
            docx_payload,
            {
                "content-type": (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                "content-disposition": "attachment; filename*=UTF-8''quality.docx",
            },
        ),
        model,
        forbidden_values=(stable_value,),
    )

    assert not any("candidate leaked" in issue for issue in issues)


def test_export_cover_uses_one_legal_h1_and_a_draft_title_subtitle():
    model, _docx_payload, _pdf_payload = _semantic_fixture()
    soup = BeautifulSoup(model.html, "html.parser")
    cover = soup.select_one("section.cover-page")

    assert cover is not None
    assert len(cover.find_all("h1")) == 1
    subtitle = cover.select_one("h2.draft-title")
    assert subtitle is not None
    assert subtitle.get_text(strip=True) == model.title


@pytest.mark.parametrize(
    "document", load_corpus().documents, ids=lambda item: item.doc_id
)
def test_every_export_cover_has_one_legal_h1(document: CorpusDocument):
    manifest = load_manifest(document.doc_id)
    assert manifest is not None
    snapshot = quality_runner._confirm_fields(
        snapshot_from_document_state(
            doc_id=document.doc_id, state={}, manifest=manifest
        ),
        manifest["fields"],
    )
    model = build_export_document(
        doc_id=document.doc_id,
        title=f"{document.doc_id}-saved-draft-title",
        manifest=manifest,
        snapshot=snapshot,
    )
    soup = BeautifulSoup(model.html, "html.parser")
    cover = soup.select_one("section.cover-page")

    assert cover is not None
    assert len(cover.find_all("h1")) == 1
    assert cover.select_one("h2.draft-title") is not None


def test_coverage_detects_noop_public_put_suite(monkeypatch):
    monkeypatch.setattr(quality_runner, "_evaluate_public_put", lambda *_args: None)

    report = run_contract_quality_evaluation(_single_document_validated())

    assert "coverage_missing" in report.invariant_errors
    assert any(
        "public_put_protection" in key
        for key in report.coverage_summary()["missing_coverage_keys"]
    )


@pytest.mark.parametrize(
    "omitted_case",
    [
        "cross-format-semantics:complete:pdf",
        "invalid-downloads:pending:pdf",
    ],
)
def test_coverage_detects_skipped_renderer_or_invalid_state(
    monkeypatch, omitted_case: str
):
    original_record = quality_runner._record

    def skip_one(report, doc_id, case_id, passed, **kwargs):
        if case_id == omitted_case:
            return None
        return original_record(report, doc_id, case_id, passed, **kwargs)

    monkeypatch.setattr(quality_runner, "_record", skip_one)

    report = run_contract_quality_evaluation(_single_document_validated())

    assert "coverage_missing" in report.invariant_errors
    assert any(
        omitted_case.rsplit(":", 1)[0] in key
        for key in report.coverage_summary()["missing_coverage_keys"]
    )


def test_coverage_detects_partial_route_suite(monkeypatch):
    def partial_suite(report, client, _get_conn, headers, document, _manifest):
        created = quality_runner._create_document(
            client, headers, document.doc_id, "partial-invalid-downloads"
        )
        quality_runner._assert_blocked_formats(
            report,
            client,
            headers,
            document.doc_id,
            created["id"],
            "missing",
            document.anchor_field,
        )

    monkeypatch.setattr(quality_runner, "_evaluate_invalid_downloads", partial_suite)

    report = run_contract_quality_evaluation(_single_document_validated())

    assert "coverage_missing" in report.invariant_errors
    assert any(
        "invalid_downloads" in key
        for key in report.coverage_summary()["missing_coverage_keys"]
    )
    assert any(
        "invalid-downloads:missing" in record.key
        for record in report.coverage_records
    )


def test_deleted_metric_family_fails_closed(monkeypatch):
    monkeypatch.setattr(quality_runner, "_evaluate_unknown_field", lambda *_args: None)

    report = run_contract_quality_evaluation(_single_document_validated())

    assert report.metric_denominators["unknown_field_acceptance_count"] == 0
    assert "zero_metric_denominator" in report.invariant_errors


@pytest.mark.contract_quality_gate
def test_contract_quality_hard_gate_is_zero_and_report_is_serializable(monkeypatch):
    def reject_network(*_args, **_kwargs):
        raise AssertionError("deterministic evaluator attempted network access")

    # Other backend tests legitimately import the chat stack. Remove those
    # cached modules for this test so the evaluator must prove it does not load
    # the LLM path itself; monkeypatch restores the original entries afterward.
    for module_name in tuple(sys.modules):
        if module_name == "app.llm" or module_name.startswith("litellm"):
            monkeypatch.delitem(sys.modules, module_name)
    monkeypatch.setattr(socket.socket, "connect", reject_network)
    report = run_contract_quality_evaluation()

    assert report.failed_cases == 0, report.failure_summary()
    assert report.passed_cases == report.total_cases
    assert all(report.metrics[name] == 0 for name in HARD_GATE_METRICS)
    encoded = json.dumps(report.to_dict(), ensure_ascii=False, sort_keys=True)
    decoded = json.loads(encoded)
    assert decoded["schema_version"] == 1
    assert len(decoded["documents"]) == 11
    assert all(value > 0 for value in decoded["metric_denominators"].values())
    coverage = decoded["coverage"]
    assert coverage["expected_count"] == coverage["actual_count"]
    assert coverage["missing_coverage_keys"] == []
    assert coverage["unexpected_duplicate_coverage_keys"] == []
    assert all(
        "cross-format-semantics:pending" in item["key"]
        or "cross-format-semantics:conflict" in item["key"]
        for item in coverage["not_applicable"]
    )
    assert {
        item["key"].split("::", 1)[0]
        for item in coverage["not_applicable"]
    } == {
        "design-partner-agreement",
        "software-license-agreement",
        "business-associate-agreement",
    }
    assert "app.llm" not in sys.modules
    assert "litellm" not in sys.modules


class _raises_corpus_error:
    def __init__(self, kind: str):
        self.kind = kind

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, traceback):
        assert exc_type is CorpusValidationError
        assert exc.kind == self.kind
        return True


def _semantic_fixture(anchor_value: str = "确定性稳定值"):
    manifest = load_manifest("mutual-nda")
    assert manifest is not None
    snapshot = snapshot_from_document_state(
        doc_id="mutual-nda",
        state={},
        manifest=manifest,
    )
    values = {
        field["key"]: (
            anchor_value
            if field["key"] == "保密用途"
            else quality_runner._valid_value(field)
        )
        for field in manifest["fields"]
    }
    snapshot = quality_runner._confirm_values(snapshot, manifest, values)
    model = build_export_document(
        doc_id="mutual-nda",
        title="确定性语义验证协议",
        manifest=manifest,
        snapshot=snapshot,
    )
    return model, render_docx(model), render_pdf(model)


def _download_response(payload: bytes, headers: dict[str, str]):
    return SimpleNamespace(status_code=200, content=payload, headers=headers)


def _docx_with_lines(lines: list[str]) -> bytes:
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _valid_report() -> QualityReport:
    report = QualityReport.empty(expected_doc_ids={"expected-doc"})
    for metric in EXPECTED_METRIC_NAMES:
        report.add(
            QualityCaseResult(
                doc_id="expected-doc",
                case_id=f"baseline-{metric}",
                passed=True,
                metric=metric,
            )
        )
    assert report.exit_code == 0
    return report


def _single_document_validated() -> ValidatedCorpus:
    validated = validate_corpus(load_corpus())
    document = validated.documents[0]
    corpus = replace(validated.corpus, documents=(document,))
    return ValidatedCorpus(
        corpus=corpus,
        catalog_doc_ids={document.doc_id},
        manifest_doc_ids={document.doc_id},
        corpus_doc_ids={document.doc_id},
    )
