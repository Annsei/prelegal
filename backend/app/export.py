"""Shared semantic rendering for server-side DOCX and PDF exports."""

from __future__ import annotations

import html
import json
import os
import re
import sys
from dataclasses import dataclass
from functools import cache
from io import BytesIO
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from markdown import markdown

from app.draft_state import (
    CONDITION_OPERATORS,
    DraftStateSnapshot,
    confirmed_field_value,
    required_field_keys,
    single_condition_matches,
)

DISCLAIMER = (
    "本文档为 AI 生成的草稿，仅供参考，不构成法律意见。"
    "正式签署前请交由执业律师审核。"
)
OPTIONAL_DEFAULT = "／（适用标准条款默认约定）"
TERM_CLASSES = {"coverpage_link", "orderform_link", "keyterms_link"}
_WHEN_RE = re.compile(r"^\s*<!--\s*when\s+(.+?)\s*-->\s*$")
_ENDWHEN_RE = re.compile(r"^\s*<!--\s*endwhen\s*-->\s*$")

_REPO_ROOT = Path(__file__).resolve().parent.parent.parent
_TEMPLATES_DIR = _REPO_ROOT / "templates"
_TEMPLATES_INDEX = _TEMPLATES_DIR / "templates.json"


class ExportTemplateError(ValueError):
    """The document's indexed template package cannot be exported."""


@dataclass(frozen=True)
class ExportField:
    """One field in the semantic verification projection for an export."""

    key: str
    label: str
    value: str | None
    rendered_value: str
    status: str
    required: bool


@dataclass(frozen=True)
class ExportConditionResult:
    """The backend-authoritative result of one manifest required_when rule."""

    field_key: str
    active: bool


@dataclass(frozen=True)
class ExportBlock:
    """An ordered text projection used to verify rendered output."""

    section: str
    order: int
    kind: str
    text: str


@dataclass(frozen=True)
class ExportDocument:
    """Authoritative HTML renderer input plus its verification projection."""

    doc_id: str
    title: str
    fields: tuple[ExportField, ...]
    condition_results: tuple[ExportConditionResult, ...]
    blocks: tuple[ExportBlock, ...]
    disclaimer: str
    html: str


@cache
def _template_index() -> dict[str, dict[str, Any]]:
    try:
        raw = json.loads(_TEMPLATES_INDEX.read_text())
    except (FileNotFoundError, json.JSONDecodeError):
        return {}
    return {
        entry["id"]: entry
        for entry in raw.get("templates", [])
        if isinstance(entry, dict) and isinstance(entry.get("id"), str)
    }


def build_export_document(
    *,
    doc_id: str,
    title: str,
    manifest: dict[str, Any],
    snapshot: DraftStateSnapshot,
) -> ExportDocument:
    """Build normalized HTML from template markdown and confirmed fields."""
    cover_markdown, terms_markdown = _load_template_markdown(doc_id)
    if not cover_markdown:
        cover_markdown = _manifest_cover_markdown(title, manifest)
    confirmed = _stable_confirmed_values(snapshot)
    active_required = set(required_field_keys(manifest, snapshot))
    cover_markdown = apply_conditional_blocks(
        cover_markdown,
        manifest=manifest,
        snapshot=snapshot,
    )
    terms_markdown = apply_conditional_blocks(
        terms_markdown,
        manifest=manifest,
        snapshot=snapshot,
    )
    lookup = _field_lookup(manifest)
    cover_html = _render_markdown_section(
        cover_markdown,
        lookup=lookup,
        confirmed=confirmed,
        cover_page=True,
    )
    cover_html = _add_draft_title_subtitle(cover_html, title)
    terms_html = _render_markdown_section(
        terms_markdown,
        lookup=lookup,
        confirmed=confirmed,
        cover_page=False,
    )
    full_html = f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{html.escape(title)}</title>
<style>{_EXPORT_CSS}</style>
</head>
<body>
<main class="export-document">
<section class="cover-page">
{cover_html}
</section>
<section class="standard-terms">{terms_html}</section>
<section class="export-disclaimer">
<h1>免责声明</h1>
<p>{DISCLAIMER}</p>
</section>
</main>
</body>
</html>"""
    fields = tuple(
        _canonical_field(field, snapshot, confirmed, active_required)
        for field in manifest.get("fields", [])
        if isinstance(field, dict) and isinstance(field.get("key"), str)
    )
    condition_results = tuple(
        ExportConditionResult(
            field_key=field["key"],
            active=field["key"] in active_required,
        )
        for field in manifest.get("fields", [])
        if isinstance(field, dict)
        and isinstance(field.get("key"), str)
        and field.get("required_when") is not None
    )
    return ExportDocument(
        doc_id=doc_id,
        title=title,
        fields=fields,
        condition_results=condition_results,
        blocks=_canonical_blocks(full_html),
        disclaimer=DISCLAIMER,
        html=full_html,
    )


def _add_draft_title_subtitle(cover_html: str, title: str) -> str:
    """Keep the legal template title as H1 and show the user's draft as metadata."""
    soup = BeautifulSoup(cover_html, "html.parser")
    legal_titles = soup.find_all("h1")
    if len(legal_titles) != 1:
        raise ExportTemplateError("Cover page must contain exactly one legal H1 title.")
    subtitle = soup.new_tag("h2", attrs={"class": "draft-title"})
    subtitle.string = title
    legal_titles[0].insert_after(subtitle)
    return str(soup)


def _stable_confirmed_values(snapshot: DraftStateSnapshot) -> dict[str, str]:
    values: dict[str, str] = {}
    for key, field in snapshot.fields.items():
        stable = field.status == "confirmed" or (
            field.status == "conflict" and field.confirmed_at is not None
        )
        if stable and isinstance(field.value, str) and field.value.strip():
            values[key] = field.value.strip()
    return values


def _canonical_field(
    field_def: dict[str, Any],
    snapshot: DraftStateSnapshot,
    confirmed: dict[str, str],
    active_required: set[str],
) -> ExportField:
    key = field_def["key"]
    label = field_def.get("label")
    display_label = label.get("zh") if isinstance(label, dict) else key
    state = snapshot.fields.get(key)
    value = confirmed.get(key)
    return ExportField(
        key=key,
        label=display_label or key,
        value=value,
        rendered_value=value or OPTIONAL_DEFAULT,
        status=state.status if state is not None else "missing",
        required=key in active_required,
    )


def _canonical_blocks(source_html: str) -> tuple[ExportBlock, ...]:
    soup = BeautifulSoup(source_html, "html.parser")
    blocks: list[ExportBlock] = []
    order = 0
    for section_name, class_name in (
        ("cover_page", "cover-page"),
        ("standard_terms", "standard-terms"),
        ("disclaimer", "export-disclaimer"),
    ):
        section = soup.find("section", class_=class_name)
        if section is None:
            continue
        for node in section.find_all(recursive=False):
            if not isinstance(node, Tag):
                continue
            projection_nodes = _projection_nodes(node)
            for projection_node in projection_nodes:
                text = " ".join(projection_node.stripped_strings)
                if not text:
                    continue
                blocks.append(
                    ExportBlock(
                        section=section_name,
                        order=order,
                        kind=projection_node.name,
                        text=text,
                    )
                )
                order += 1
    return tuple(blocks)


def _projection_nodes(node: Tag) -> list[Tag]:
    atomic = {"h1", "h2", "h3", "h4", "h5", "h6", "p"}
    if node.name in atomic:
        return [node]
    if node.name == "table":
        return node.find_all(["th", "td"])
    if node.name not in {"ol", "ul", "li", "div", "section", "blockquote"}:
        return []
    projected: list[Tag] = []
    for child in node.children:
        if isinstance(child, Tag):
            projected.extend(_projection_nodes(child))
    return projected or [node]


def apply_conditional_blocks(
    source: str,
    *,
    manifest: dict[str, Any],
    snapshot: DraftStateSnapshot,
) -> str:
    """Resolve non-nested markdown condition blocks using kernel semantics."""
    manifest_keys = {
        field["key"]
        for field in manifest.get("fields", [])
        if isinstance(field, dict) and isinstance(field.get("key"), str)
    }
    output: list[str] = []
    active: tuple[bool, int] | None = None

    for line_number, line in enumerate(source.splitlines(keepends=True), start=1):
        marker = line.rstrip("\r\n")
        when_match = _WHEN_RE.fullmatch(marker)
        if when_match:
            if active is not None:
                raise ExportTemplateError(
                    f"Conditional blocks cannot be nested (line {line_number})."
                )
            condition = _parse_conditional_marker(
                when_match.group(1),
                manifest_keys=manifest_keys,
                line_number=line_number,
            )
            stable_value = confirmed_field_value(snapshot, condition["field"])
            include = stable_value is None or single_condition_matches(
                condition,
                snapshot,
            )
            active = (include, line_number)
            continue

        if _ENDWHEN_RE.fullmatch(marker):
            if active is None:
                raise ExportTemplateError(
                    "Conditional end marker has no opening marker "
                    f"(line {line_number})."
                )
            active = None
            continue

        stripped = marker.strip()
        if stripped.startswith("<!-- when") or stripped.startswith("<!-- endwhen"):
            raise ExportTemplateError(
                f"Malformed conditional marker at line {line_number}."
            )
        if active is None or active[0]:
            output.append(line)

    if active is not None:
        raise ExportTemplateError(
            f"Conditional block opened at line {active[1]} is not closed."
        )
    return "".join(output)


def _parse_conditional_marker(
    payload: str,
    *,
    manifest_keys: set[str],
    line_number: int,
) -> dict[str, Any]:
    try:
        condition = json.loads(payload)
    except json.JSONDecodeError as exc:
        raise ExportTemplateError(
            f"Conditional marker at line {line_number} contains invalid JSON."
        ) from exc
    if not isinstance(condition, dict):
        raise ExportTemplateError(
            f"Conditional marker at line {line_number} must contain an object."
        )
    field_key = condition.get("field")
    if (
        not isinstance(field_key, str)
        or not field_key.strip()
        or field_key not in manifest_keys
    ):
        raise ExportTemplateError(
            f"Conditional marker at line {line_number} references unknown "
            "manifest field."
        )
    op = condition["op"] if "op" in condition else "equals"
    if not isinstance(op, str) or op not in CONDITION_OPERATORS:
        raise ExportTemplateError(
            f"Conditional marker at line {line_number} uses unsupported operator."
        )
    if op in {"equals", "not_equals"} and not isinstance(
        condition.get("value"), str
    ):
        raise ExportTemplateError(
            f"Conditional marker at line {line_number} requires a string value."
        )
    values = condition.get("values")
    if op == "in" and (
        not isinstance(values, list)
        or not values
        or not all(isinstance(value, str) for value in values)
    ):
        raise ExportTemplateError(
            f"Conditional marker at line {line_number} requires string values."
        )
    return condition


def render_docx(model: ExportDocument) -> bytes:
    """Render the normalized HTML model to an in-memory DOCX file."""
    document = Document()
    _configure_docx(document)
    soup = BeautifulSoup(model.html, "html.parser")
    main = soup.find("main")
    if main is None:
        raise ExportTemplateError("Normalized export model has no document body.")
    _render_block_children(document, main)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_pdf(model: ExportDocument) -> bytes:
    """Render the normalized HTML model to PDF with WeasyPrint."""
    # Homebrew supplies the same Pango libraries used by the Docker runtime,
    # but macOS does not search /opt/homebrew/lib for dlopen by default.
    if sys.platform == "darwin" and Path("/opt/homebrew/lib").is_dir():
        os.environ.setdefault("DYLD_FALLBACK_LIBRARY_PATH", "/opt/homebrew/lib")
    from weasyprint import HTML

    return HTML(string=model.html, base_url=str(_TEMPLATES_DIR)).write_pdf()


def _load_template_markdown(doc_id: str) -> tuple[str | None, str]:
    entry = _template_index().get(doc_id)
    if entry is None:
        raise ExportTemplateError(f"Unknown template package: {doc_id}")
    content: dict[str, str] = {}
    for file_entry in entry.get("files", []):
        if not isinstance(file_entry, dict):
            continue
        file_type = file_entry.get("type")
        rel_path = file_entry.get("path")
        if file_type not in {"cover_page", "standard_terms"} or not isinstance(
            rel_path, str
        ):
            continue
        try:
            path = (_TEMPLATES_DIR / rel_path).resolve()
            path.relative_to(_TEMPLATES_DIR.resolve())
            content[file_type] = path.read_text()
        except (FileNotFoundError, ValueError):
            continue
    if not content.get("standard_terms"):
        raise ExportTemplateError(
            f"Template package {doc_id} requires standard_terms."
        )
    return content.get("cover_page"), content["standard_terms"]


def _manifest_cover_markdown(title: str, manifest: dict[str, Any]) -> str:
    """Generate the same structured cover-page semantics used by preview."""
    lines = [
        f"# {title} · 封面页",
        "",
        "本封面页与标准条款共同构成双方协议；确认内容对应标准条款中的同名字段。",
    ]
    fields = [field for field in manifest.get("fields", []) if isinstance(field, dict)]
    for section in manifest.get("sections", []):
        if not isinstance(section, dict):
            continue
        section_key = section.get("key")
        section_fields = [
            field for field in fields if field.get("section") == section_key
        ]
        if not section_fields:
            continue
        label = section.get("label")
        section_title = label.get("zh") if isinstance(label, dict) else section_key
        lines.extend(
            [
                "",
                f"## {section_title}",
                "",
                "| 项目 | 约定内容 |",
                "| --- | --- |",
            ]
        )
        for field in section_fields:
            field_label = field.get("label")
            display = (
                field_label.get("zh")
                if isinstance(field_label, dict)
                else field.get("key")
            )
            key = field.get("key")
            lines.append(
                f'| {display} | <span class="coverpage_link">{key}</span> |'
            )
    return "\n".join(lines)


def _field_lookup(manifest: dict[str, Any]) -> dict[str, dict[str, Any]]:
    lookup: dict[str, dict[str, Any]] = {}
    for field in manifest.get("fields", []):
        if not isinstance(field, dict) or not isinstance(field.get("key"), str):
            continue
        lookup[field["key"]] = field
        for alias in field.get("aliases", []):
            if isinstance(alias, str):
                lookup[alias] = field
    return lookup


def _render_markdown_section(
    source: str,
    *,
    lookup: dict[str, dict[str, Any]],
    confirmed: dict[str, str],
    cover_page: bool,
) -> str:
    rendered = markdown(source, extensions=["tables", "sane_lists"])
    soup = BeautifulSoup(rendered, "html.parser")
    for span in soup.find_all("span"):
        classes = set(span.get("class", []))
        if not classes.intersection(TERM_CLASSES):
            continue
        reference = span.get_text(strip=True)
        field = lookup.get(reference)
        if field is None:
            continue
        value = confirmed.get(field["key"])
        if value:
            span.clear()
            span.append(value)
            classes.add("filled-term")
        elif cover_page:
            span.clear()
            span.append(OPTIONAL_DEFAULT)
            classes.add("default-term")
        else:
            classes.add("unfilled-term")
        span["class"] = sorted(classes)
    return str(soup)


def _configure_docx(document: DocxDocument) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for level in range(1, 7):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    document.styles["Heading 1"].font.size = Pt(18)
    document.styles["Heading 2"].font.size = Pt(14)
    document.styles["Heading 3"].font.size = Pt(12)


def _render_block_children(container: DocxDocument, parent: Tag) -> None:
    for child in parent.children:
        if not isinstance(child, Tag):
            continue
        _render_block(container, child)


def _render_block(document: DocxDocument, node: Tag) -> None:
    name = node.name.lower()
    if name == "section":
        if "export-disclaimer" in node.get("class", []):
            document.add_page_break()
        _render_block_children(document, node)
        return
    if re.fullmatch(r"h[1-6]", name):
        paragraph = document.add_heading(level=int(name[1]))
        _append_inline(paragraph, node)
        if name == "h1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    if name == "p":
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.35
        _append_inline(paragraph, node)
        return
    if name in {"ol", "ul"}:
        _render_list(document, node, level=0)
        return
    if name == "table":
        _render_table(document, node)
        return
    if name == "blockquote":
        paragraph = document.add_paragraph(style="Intense Quote")
        _append_inline(paragraph, node)
        return
    if name == "hr":
        document.add_paragraph("—" * 24)
        return
    _render_block_children(document, node)


def _render_list(document: DocxDocument, node: Tag, *, level: int) -> None:
    ordered = node.name.lower() == "ol"
    style_base = "List Number" if ordered else "List Bullet"
    style_name = style_base if level == 0 else f"{style_base} {min(level + 1, 3)}"
    for item in node.find_all("li", recursive=False):
        paragraph = document.add_paragraph(style=style_name)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.25
        for child in item.children:
            if isinstance(child, Tag) and child.name.lower() in {"ol", "ul"}:
                _render_list(document, child, level=level + 1)
            else:
                _append_inline_node(paragraph, child)


def _render_table(document: DocxDocument, node: Tag) -> None:
    rows = node.find_all("tr")
    if not rows:
        return
    column_count = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
    table = document.add_table(rows=len(rows), cols=max(column_count, 1))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        cells = row.find_all(["th", "td"], recursive=False)
        for column_index, source_cell in enumerate(cells):
            target = table.cell(row_index, column_index)
            paragraph = target.paragraphs[0]
            _append_inline(paragraph, source_cell)
            if source_cell.name.lower() == "th":
                for run in paragraph.runs:
                    run.bold = True


def _append_inline(paragraph, parent: Tag) -> None:
    for child in parent.children:
        _append_inline_node(paragraph, child)


def _append_inline_node(
    paragraph,
    node,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    if isinstance(node, NavigableString):
        if str(node):
            _add_run(paragraph, str(node), bold=bold, italic=italic)
        return
    if not isinstance(node, Tag):
        return
    name = node.name.lower()
    if name == "br":
        paragraph.add_run().add_break()
        return
    next_bold = bold or name in {"strong", "b"}
    next_italic = italic or name in {"em", "i"}
    filled = name == "span" and "filled-term" in node.get("class", [])
    defaulted = name == "span" and "default-term" in node.get("class", [])
    if name in {"ol", "ul"}:
        return
    if name == "code":
        run = _add_run(paragraph, node.get_text(), bold=next_bold, italic=next_italic)
        run.font.name = "Courier New"
        return
    if filled or defaulted:
        run = _add_run(
            paragraph,
            node.get_text(),
            bold=next_bold or filled,
            italic=next_italic,
        )
        run.underline = True
        return
    for child in node.children:
        _append_inline_node(
            paragraph,
            child,
            bold=next_bold,
            italic=next_italic,
        )


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    return run


_EXPORT_CSS = """
@page { size: A4; margin: 22mm 22mm 24mm; }
body {
  color: #111;
  font-family: "Noto Serif CJK SC", "Source Han Serif SC", "Songti SC",
    "STSong", "SimSun", serif;
  font-size: 10.5pt;
  line-height: 1.6;
}
h1, h2, h3, h4, h5, h6 {
  color: #111;
  font-family: "Noto Sans CJK SC", "Source Han Sans SC", "PingFang SC",
    "Microsoft YaHei", sans-serif;
  break-after: avoid;
}
h1 { font-size: 18pt; text-align: center; margin: 0 0 18pt; }
h2 { font-size: 14pt; margin: 16pt 0 8pt; }
.draft-title {
  color: #444;
  font-size: 10.5pt;
  font-weight: 400;
  margin: -10pt 0 18pt;
  text-align: center;
}
h3 { font-size: 12pt; margin: 12pt 0 6pt; }
p { margin: 0 0 7pt; }
ol, ul { margin: 5pt 0 8pt 20pt; padding-left: 12pt; }
li { margin: 0 0 4pt; }
table { width: 100%; border-collapse: collapse; margin: 8pt 0 14pt; }
th, td { border: 0.7pt solid #666; padding: 5pt 6pt; vertical-align: top; }
th { font-family: "Noto Sans CJK SC", "PingFang SC", sans-serif; background: #f2f2f2; }
.filled-term { font-weight: 700; text-decoration: underline; }
.default-term { color: #444; text-decoration: underline; }
.unfilled-term { font-weight: 600; }
.standard-terms { break-before: page; }
.export-disclaimer { break-before: page; }
.export-disclaimer p { font-size: 11pt; }
"""
