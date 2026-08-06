"""Server-side DOCX/PDF export contract tests."""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfReader

from app.draft_state import DraftStateSnapshot, FieldState
from app.export import ExportTemplateError, build_export_document, render_docx
from app.manifests import load_manifest

CATALOG_DOC_IDS = (
    "mutual-nda",
    "cloud-service-agreement",
    "design-partner-agreement",
    "service-level-agreement",
    "professional-services-agreement",
    "data-processing-agreement",
    "software-license-agreement",
    "partnership-agreement",
    "pilot-agreement",
    "business-associate-agreement",
    "ai-addendum",
)

DISCLAIMER = (
    "本文档为 AI 生成的草稿，仅供参考，不构成法律意见。"
    "正式签署前请交由执业律师审核。"
)
OPTIONAL_DEFAULT = "／（适用标准条款默认约定）"


def _register(client, email: str) -> dict[str, str]:
    response = client.post(
        "/api/auth/register",
        json={"email": email, "password": "export-test-password"},
    )
    assert response.status_code == 201
    return {"Authorization": f"Bearer {response.json()['token']}"}


def _create_document(
    client,
    headers: dict[str, str],
    doc_id: str,
    *,
    title: str = "导出测试协议",
) -> dict:
    response = client.post(
        "/api/documents",
        headers=headers,
        json={"doc_id": doc_id, "title": title, "state": {}},
    )
    assert response.status_code == 201
    return response.json()


def _field_value(field: dict, overrides: dict[str, str] | None = None) -> str:
    if overrides and field["key"] in overrides:
        return overrides[field["key"]]
    if field["type"] == "date":
        return "2026-08-03"
    return "已确认导出测试值"


def _confirm_required_fields(
    client,
    headers: dict[str, str],
    document: dict,
    *,
    overrides: dict[str, str] | None = None,
) -> dict:
    manifest = load_manifest(document["doc_id"])
    assert manifest is not None
    operations = [
        {
            "op": "confirm",
            "key": field["key"],
            "value": _field_value(field, overrides),
        }
        for field in manifest["fields"]
        if field["required"]
    ]
    response = client.post(
        f"/api/documents/{document['id']}/field-patches",
        headers=headers,
        json={
            "patch_id": f"export-required-{document['doc_id']}",
            "base_revision": 0,
            "source": "form",
            "operations": operations,
        },
    )
    assert response.status_code == 200, response.text
    return response.json()["snapshot"]


def _docx_text(payload: bytes) -> str:
    document = Document(BytesIO(payload))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def _pilot_snapshot(
    pricing_status: str,
    pricing_value: str | None,
) -> DraftStateSnapshot:
    confirmed_at = (
        "2026-08-06T00:00:00+00:00" if pricing_status == "confirmed" else None
    )
    return DraftStateSnapshot(
        doc_id="pilot-agreement",
        fields={
            "试点收费方式": FieldState(
                key="试点收费方式",
                status=pricing_status,
                value=pricing_value,
                confirmed_at=confirmed_at,
                confirmed_by_user_id=1 if confirmed_at else None,
            )
        },
    )


@pytest.mark.parametrize("doc_id", CATALOG_DOC_IDS)
def test_all_catalog_documents_export_valid_docx(client, doc_id: str):
    headers = _register(client, f"export-{doc_id}@example.com")
    document = _create_document(client, headers, doc_id)
    _confirm_required_fields(client, headers, document)

    response = client.get(
        f"/api/documents/{document['id']}/download?format=docx",
        headers=headers,
    )

    assert response.status_code == 200, response.text
    assert response.content.startswith(b"PK")
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "filename*=UTF-8''" in response.headers["content-disposition"]
    assert DISCLAIMER in _docx_text(response.content)


def test_docx_uses_confirmed_values_and_optional_default_without_pending_value(
    client,
):
    headers = _register(client, "docx-semantics@example.com")
    document = _create_document(
        client,
        headers,
        "service-level-agreement",
        title="中文服务等级协议",
    )
    snapshot = _confirm_required_fields(
        client,
        headers,
        document,
        overrides={"客户": "导出测试客户有限公司"},
    )
    pending_value = "不得出现在导出文件中的候选值"
    proposed = client.post(
        f"/api/documents/{document['id']}/field-patches",
        headers=headers,
        json={
            "patch_id": "export-pending-optional",
            "base_revision": snapshot["revision"],
            "source": "llm",
            "operations": [
                {
                    "op": "propose",
                    "key": "故障响应与修复时限",
                    "value": pending_value,
                },
            ],
        },
    )
    assert proposed.status_code == 200

    response = client.get(
        f"/api/documents/{document['id']}/download?format=docx",
        headers=headers,
    )

    assert response.status_code == 200
    text = _docx_text(response.content)
    assert text.count("导出测试客户有限公司") >= 2
    assert OPTIONAL_DEFAULT in text
    assert pending_value not in text
    disposition = response.headers["content-disposition"]
    encoded_title = (
        "%E4%B8%AD%E6%96%87%E6%9C%8D%E5%8A%A1"
        "%E7%AD%89%E7%BA%A7%E5%8D%8F%E8%AE%AE"
    )
    assert encoded_title in disposition


def test_pdf_is_valid_and_contains_confirmed_chinese_text(client):
    headers = _register(client, "pdf-export@example.com")
    document = _create_document(client, headers, "service-level-agreement")
    _confirm_required_fields(
        client,
        headers,
        document,
        overrides={"客户": "中文PDF测试客户有限公司"},
    )

    response = client.get(
        f"/api/documents/{document['id']}/download?format=pdf",
        headers=headers,
    )

    assert response.status_code == 200
    assert response.content.startswith(b"%PDF")
    reader = PdfReader(BytesIO(response.content))
    assert len(reader.pages) > 0
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "中文PDF测试客户有限公司" in text
    assert DISCLAIMER in text


def test_download_requires_authentication(client):
    response = client.get("/api/documents/1/download?format=docx")
    assert response.status_code == 401


def test_download_hides_other_users_document(client):
    owner = _register(client, "export-owner@example.com")
    outsider = _register(client, "export-outsider@example.com")
    document = _create_document(client, owner, "mutual-nda")

    response = client.get(
        f"/api/documents/{document['id']}/download?format=docx",
        headers=outsider,
    )

    assert response.status_code == 404


def test_download_rejects_unknown_format(client):
    headers = _register(client, "bad-export-format@example.com")
    document = _create_document(client, headers, "mutual-nda")

    response = client.get(
        f"/api/documents/{document['id']}/download?format=txt",
        headers=headers,
    )

    assert response.status_code == 422


def test_download_rejects_unmanaged_document_with_stable_409(client):
    headers = _register(client, "unmanaged-export@example.com")
    document = _create_document(client, headers, "unmanaged-test-document")

    response = client.get(
        f"/api/documents/{document['id']}/download?format=docx",
        headers=headers,
    )

    assert response.status_code == 409
    assert response.json()["detail"]["validation_errors"][0]["kind"] == (
        "unsupported_document_schema"
    )


def test_download_returns_readiness_shape_for_missing_required_fields(client):
    headers = _register(client, "blocked-export@example.com")
    document = _create_document(client, headers, "mutual-nda")

    response = client.get(
        f"/api/documents/{document['id']}/download?format=docx",
        headers=headers,
    )

    assert response.status_code == 409
    detail = response.json()["detail"]
    assert detail["validation_errors"][0]["kind"] == "download_blocked"
    assert detail["unresolved_required_fields"]


def test_download_blocks_active_conditional_required_fields(client):
    headers = _register(client, "conditional-export@example.com")
    document = _create_document(client, headers, "pilot-agreement")
    _confirm_required_fields(
        client,
        headers,
        document,
        overrides={"试点收费方式": "付费"},
    )

    response = client.get(
        f"/api/documents/{document['id']}/download?format=docx",
        headers=headers,
    )

    assert response.status_code == 409
    assert set(response.json()["detail"]["unresolved_required_fields"]) == {
        "试点费用",
        "付款安排",
    }


@pytest.mark.parametrize(
    ("status", "value", "paid_terms_visible"),
    [
        ("confirmed", "付费", True),
        ("confirmed", "免费", False),
        ("pending_confirmation", "免费", True),
    ],
)
def test_pilot_docx_conditionally_renders_paid_terms(
    status: str,
    value: str,
    paid_terms_visible: bool,
):
    manifest = load_manifest("pilot-agreement")
    assert manifest is not None
    model = build_export_document(
        doc_id="pilot-agreement",
        title="试点协议",
        manifest=manifest,
        snapshot=_pilot_snapshot(status, value),
    )

    text = _docx_text(render_docx(model))

    assert ("逾期支付试点费用" in text) is paid_terms_visible
    assert ("退还已预付的试点费用" in text) is paid_terms_visible


@pytest.mark.parametrize(
    ("source", "message"),
    [
        (
            '<!-- when {"field":"试点收费方式","op":"equals","value":"付费"} -->\n'
            "未闭合条款",
            "not closed",
        ),
        (
            '<!-- when {"field":"试点收费方式","op":"equals","value":"付费"} -->\n'
            '<!-- when {"field":"试点收费方式","op":"equals","value":"免费"} -->\n'
            "嵌套条款\n<!-- endwhen -->\n<!-- endwhen -->",
            "nested",
        ),
        (
            '<!-- when {"field":"不存在字段","op":"equals","value":"付费"} -->\n'
            "未知字段条款\n<!-- endwhen -->",
            "unknown manifest field",
        ),
        ("<!-- endwhen -->\n无起始标记", "no opening marker"),
    ],
)
def test_export_rejects_invalid_conditional_blocks(
    monkeypatch: pytest.MonkeyPatch,
    source: str,
    message: str,
):
    manifest = load_manifest("pilot-agreement")
    assert manifest is not None
    monkeypatch.setattr(
        "app.export._load_template_markdown",
        lambda _doc_id: (None, source),
    )

    with pytest.raises(ExportTemplateError, match=message):
        build_export_document(
            doc_id="pilot-agreement",
            title="试点协议",
            manifest=manifest,
            snapshot=_pilot_snapshot("confirmed", "付费"),
        )


def test_export_supports_multiple_blocks_and_required_when_operators(
    monkeypatch: pytest.MonkeyPatch,
):
    source = """<!-- when {"field":"试点收费方式","op":"equals","value":"付费"} -->
等于条件条款
<!-- endwhen -->
<!-- when {"field":"试点收费方式","op":"not_equals","value":"免费"} -->
不等于条件条款
<!-- endwhen -->
<!-- when {"field":"试点收费方式","op":"in","values":["付费","内部试点"]} -->
集合条件条款
<!-- endwhen -->
"""
    manifest = load_manifest("pilot-agreement")
    assert manifest is not None
    monkeypatch.setattr(
        "app.export._load_template_markdown",
        lambda _doc_id: (None, source),
    )

    model = build_export_document(
        doc_id="pilot-agreement",
        title="试点协议",
        manifest=manifest,
        snapshot=_pilot_snapshot("confirmed", "付费"),
    )

    assert "等于条件条款" in model.html
    assert "不等于条件条款" in model.html
    assert "集合条件条款" in model.html
    assert "<!-- when" not in model.html
